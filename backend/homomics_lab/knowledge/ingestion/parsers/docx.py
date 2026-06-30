"""DOCX parser using python-docx (optional dependency)."""

from pathlib import Path

from homomics_lab.knowledge.ingestion.models import DocumentSource, ParsedDocument
from homomics_lab.knowledge.ingestion.parsers.base import DocumentParser


class DOCXParser(DocumentParser):
    """Extract text from Microsoft Word documents."""

    def supports(self, mime_type: str, extension: str) -> bool:
        return extension == "docx" or mime_type in {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        }

    async def parse(self, source: DocumentSource) -> ParsedDocument:
        try:
            import docx
        except ImportError as exc:
            raise RuntimeError(
                "DOCX parsing requires 'python-docx'. Install with: uv pip install python-docx "
                "or add the 'knowledge' optional dependency group."
            ) from exc

        path = Path(source.source)
        document = docx.Document(str(path))
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        tables = []
        for table in document.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append(rows)

        title = None
        if paragraphs:
            first = paragraphs[0]
            if len(first) < 200:
                title = first

        return ParsedDocument(
            source=source,
            title=title,
            text=full_text,
            pages=[full_text],
            paragraphs=paragraphs,
            metadata={
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "tables": tables,
            },
        )
